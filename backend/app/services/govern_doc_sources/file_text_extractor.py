"""Extract markdown-ish text from uploaded PDF/DOCX/XLSX Knowledge Doc files.

The public contract is intentionally small: never raise, always return
{ok, text, error}. Callers decide whether to store the source blob.
"""
from __future__ import annotations

import io
import logging
from zipfile import BadZipFile

logger = logging.getLogger(__name__)

_MAX_XLSX_ROWS_PER_SHEET = 500
_MAX_XLSX_COLS_PER_SHEET = 50


def _md_cell(value) -> str:
    text = "" if value is None else str(value)
    return text.replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br>").replace("|", "\\|").strip()


def _markdown_table(rows: list[list[str]]) -> str:
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    out = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    out.extend("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join(out)


def ocr_available() -> bool:
    """Whether local OCR can run. Reported rather than assumed, because a missing
    tesseract binary makes scanned documents silently index as empty."""
    try:
        import pypdfium2  # noqa: F401
        import pytesseract

        pytesseract.get_tesseract_version()
        return True
    except Exception:  # noqa: BLE001
        return False


def _ocr_pdf_page(pdf, page_index: int) -> str:
    """Read one page as an image, LOCALLY.

    Local by design. A cloud OCR service would send the page IMAGE — everything
    visible on the page, not just its prose — to a third party, and that is what
    `external_processing = 'full'` exists to authorise. Tesseract here means a
    scanned document becomes searchable with nothing leaving the building, so no
    permission is required and none is claimed.

    Vietnamese first in the language list, with English as a fallback for the
    mixed-language documents this corpus is full of. Without `vie`, diacritics
    come back mangled — and mangled text does not fail loudly, it just quietly
    becomes an index nobody can search.
    """
    import pytesseract

    page = pdf[page_index]
    bitmap = page.render(scale=_OCR_SCALE)
    try:
        image = bitmap.to_pil()
    finally:
        close = getattr(bitmap, "close", None)
        if close:
            close()
    return (pytesseract.image_to_string(image, lang="vie+eng") or "").strip()


def _ocr_pdf_pages(data: bytes, needed: set[int]) -> dict[int, str]:
    """OCR only the pages that produced no text. Rendering and reading a page is
    expensive; doing it for pages that already extracted cleanly would multiply
    the cost of every ordinary PDF for no gain."""
    if not needed:
        return {}
    if not ocr_available():
        logger.warning(
            "file_text_extractor: %s page(s) have no extractable text and local OCR "
            "is unavailable — those pages will not be indexed", len(needed),
        )
        return {}
    out: dict[int, str] = {}
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(io.BytesIO(data))
        try:
            for index in sorted(needed):
                if index - 1 >= len(pdf):
                    continue
                try:
                    text = _ocr_pdf_page(pdf, index - 1)
                except Exception:  # noqa: BLE001 — one bad page must not lose the rest
                    logger.warning("file_text_extractor: OCR failed on page %s", index, exc_info=True)
                    continue
                if len(text) >= _OCR_TEXT_FLOOR:
                    out[index] = text
        finally:
            pdf.close()
    except Exception:  # noqa: BLE001
        logger.warning("file_text_extractor: OCR pass failed", exc_info=True)
    if out:
        logger.info("file_text_extractor: OCR recovered %s scanned page(s)", len(out))
    return out


#: A page yielding fewer than this many characters of extractable text is treated
#: as an IMAGE of a page rather than a page of text. Scanned pages are not empty —
#: they usually carry a few stray glyphs from a letterhead or a page number — so a
#: strict zero would miss most of them.
_OCR_TEXT_FLOOR = 24

#: Render scale. Tesseract wants roughly 300 DPI; PDF user space is 72 DPI.
_OCR_SCALE = 300 / 72


def _pdf_structured(data: bytes) -> tuple[list[dict], set[int]]:
    """A PDF as BLOCKS — text, tables and figures, each with page and geometry.

    Replaces the markdown-first path. Flattening to text first destroyed the
    column layout, the reading order, the figure regions and the coordinates a
    citation could point at, none of which can be recovered afterwards. The
    markdown is still produced (see `blocks_to_markdown`) because the editor and
    the version snapshots are markdown, but it is now derived FROM the blocks
    rather than being the only thing that survives.

    Returns `(blocks, pages_with_no_text)`; the second is what OCR is asked for.
    """
    import pdfplumber

    from app.services.govern_doc_sources.pdf_layout import page_blocks

    blocks: list[dict] = []
    empty_pages: set[int] = set()
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            try:
                found = page.find_tables()
            except Exception:  # noqa: BLE001
                found = []
            table_bboxes = [t.bbox for t in found]

            page_items: list[dict] = []
            for table in found:
                try:
                    rows = table.extract() or []
                except Exception:  # noqa: BLE001
                    continue
                cells = [[_md_cell(c) for c in (row or [])] for row in rows]
                markdown = _markdown_table(cells)
                if not markdown:
                    continue
                header = "\n".join(markdown.split("\n")[:2])
                page_items.append({
                    "kind": "table", "text": markdown, "page": index,
                    "bbox": list(table.bbox), "table_header": header,
                    "meta": {"rows": len(cells)},
                })

            for block in page_blocks(page, index, table_bboxes):
                page_items.append({
                    "kind": block.kind, "text": block.text, "page": index,
                    "bbox": list(block.bbox), "table_header": block.table_header,
                    "meta": {**block.meta, "column": block.column},
                })

            prose = "".join(
                item["text"] for item in page_items if item["kind"] != "figure"
            ).strip()
            if len(prose) < _OCR_TEXT_FLOOR:
                # An image of a page. Kept as an OCR candidate rather than dropped:
                # silently skipping it is how a scanned document became an empty
                # index. Figures found on it are still recorded.
                empty_pages.add(index)
                page_items = [i for i in page_items if i["kind"] == "figure"]
            blocks.extend(page_items)
    return blocks, empty_pages


def blocks_to_markdown(blocks: list[dict]) -> str:
    """Blocks back to markdown, for the editor and the version snapshot.

    `## Page N` headings are emitted because the rest of the pipeline reads page
    numbers from them for non-PDF sources, and because a person opening the editor
    expects to see where the pages were.
    """
    parts: list[str] = []
    current_page: int | None = None
    for block in blocks:
        page = block.get("page")
        if page and page != current_page:
            parts.append("## Page %d" % page)
            current_page = page
        text = (block.get("text") or "").strip()
        if not text:
            continue
        if block.get("kind") == "figure":
            parts.append("> [hình] " + text)
        else:
            parts.append(text)
    return "\n\n".join(parts).strip()


def extract_pdf_blocks(data: bytes) -> list[dict]:
    """The structured extraction, with OCR filled in for scanned pages."""
    blocks, empty_pages = _pdf_structured(data)
    for page, text in _ocr_pdf_pages(data, empty_pages).items():
        blocks.append({
            "kind": "paragraph", "text": text, "page": page,
            "bbox": None, "table_header": None,
            "meta": {"ocr": True, "column": 0},
        })
    blocks.sort(key=lambda b: (
        b.get("page") or 0,
        (b.get("meta") or {}).get("column", 0),
        (b.get("bbox") or [0, 0, 0, 0])[1],
    ))
    return blocks


def _extract_pdf(data: bytes) -> str:
    """Prefer layout-aware extraction; fall back to a plain text read.

    pdfplumber is slower and can choke on unusual PDFs, so pypdf remains the
    safety net — a document that yields flat text is still far better than an
    upload that fails outright.
    """
    try:
        blocks = extract_pdf_blocks(data)
        if blocks:
            return blocks_to_markdown(blocks)
    except Exception:  # noqa: BLE001
        logger.warning("file_text_extractor: structured extraction failed, "
                       "falling back to pypdf", exc_info=True)

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data), strict=False)
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            logger.info("file_text_extractor: encrypted PDF could not be opened with an empty password")

    pages = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:  # noqa: BLE001
            logger.warning("file_text_extractor: failed extracting PDF page %s", index, exc_info=True)
            continue
        if text:
            pages.append(f"## Page {index}\n\n{text}")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (getattr(paragraph.style, "name", "") or "").lower()
        if style.startswith("heading"):
            digits = "".join(ch for ch in style if ch.isdigit())
            level = max(1, min(6, int(digits or "2")))
            blocks.append(f"{'#' * level} {text}")
        elif style.startswith("list"):
            blocks.append(f"- {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        table_md = _markdown_table([[_md_cell(cell.text) for cell in row.cells] for row in table.rows])
        if table_md:
            blocks.append(table_md)

    images: list[str] = []
    for rel in document.part.rels.values():
        if "image" in getattr(rel, "reltype", ""):
            images.append(getattr(rel, "target_ref", "") or getattr(rel, "rId", "image"))
    if images:
        blocks.append("## Embedded images\n\n" + "\n".join(f"- {name}" for name in images))

    return "\n\n".join(blocks)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sections: list[str] = []
    for ws in wb.worksheets:
        rows: list[list[str]] = []
        for index, row in enumerate(ws.iter_rows(values_only=True)):
            if index >= _MAX_XLSX_ROWS_PER_SHEET:
                rows.append([f"... ({ws.max_row} rows total, truncated)"])
                break
            cells = [_md_cell(value) for value in row[:_MAX_XLSX_COLS_PER_SHEET]]
            if any(cells):
                rows.append(cells)
        table_md = _markdown_table(rows)
        if table_md:
            sections.append(f"## {ws.title}\n\n{table_md}")
    return "\n\n".join(sections)


_EXTRACTORS = {".pdf": _extract_pdf, ".docx": _extract_docx, ".xlsx": _extract_xlsx}


def extract_text(data: bytes, filename: str) -> dict:
    """Returns {ok, text, error}."""
    ext = "." + (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else ""
    extractor = _EXTRACTORS.get(ext)
    if not extractor:
        return {"ok": False, "error": f"Unsupported file type '{ext or filename}'. Allowed: .pdf, .docx, .xlsx"}
    try:
        text = (extractor(data) or "").strip()
    except BadZipFile:
        logger.warning("file_text_extractor: invalid Office zip for %s", filename, exc_info=True)
        return {
            "ok": False,
            "error": f"{ext} file is not a valid Office Open XML file. If this is an older .xls file, save it as .xlsx and upload again.",
        }
    except Exception as exc:  # noqa: BLE001
        logger.warning("file_text_extractor: failed extracting %s", filename, exc_info=True)
        return {"ok": False, "error": f"Failed to read {ext} file: {exc}"}
    if not text:
        return {"ok": False, "error": "No readable text found in this file."}
    return {"ok": True, "text": text}
