"""Govern Knowledge Doc — external sources + configurable embedding.

Covers the pure-function surface added for the Source/Embedding tabs:
  - chunk_doc() parity with the original hardcoded paragraph algorithm
    (regression guard — nothing calling chunk_doc(body) with no kwargs may
    ever change behavior).
  - heading/fixed strategies and overlap actually do something sensible.
  - preview_chunks() is byte-identical to chunk_doc() for the same args (the
    whole point of "preview always matches what will be embedded").
  - file_text_extractor round-trips a real (tiny) DOCX/XLSX/PDF fixture.
  - google_doc_fetcher._flatten() turns a Docs API structural document into
    the expected markdown-ish text.
"""
from __future__ import annotations

#: Built from the retriever's OWN column list, so a fixture cannot describe a row
#: shape production does not produce. Two fixtures hard-coded a fifteen-tuple and
#: broke with `IndexError` when the SELECT grew — pointing at neither the SELECT
#: nor the reader.
from app.services.dashboard_ai_bot.govern_doc_embeddings import (  # noqa: E402
    CHUNK_HYDRATION_COLUMNS,
)


def chunk_row(**values):
    """One hydration row, defaults for everything not named."""
    defaults = {
        "trust": "authored", "chunk_index": 0, "page": None,
        "model_version": "text-embedding-3-small", "heading_path": None,
        "block_kind": "paragraph", "token_count": 4, "section_index": 0,
        "block_from": 0, "block_to": 0, "source_version": 0,
        "last_verified_at": None, "review_date": None, "importance": "normal",
        "sensitivity": "internal", "owner": None, "status": "Published",
        "updated_at": None, "doc_type": "article",
    }
    defaults.update(values)
    missing = [n for n in CHUNK_HYDRATION_COLUMNS if n not in defaults]
    assert not missing, "chunk_row() chua biet cot moi: %s" % ", ".join(missing)
    return tuple(defaults[name] for name in CHUNK_HYDRATION_COLUMNS)

import os

os.environ.setdefault("DATABASE_URL", "sqlite:///./test_govern_doc_sources.db")
os.environ.setdefault("DATA_DIR", ".testdata")

import io

import pytest

from app.services.dashboard_ai_bot.govern_doc_embeddings import (
    chunk_doc,
    preview_chunks,
    _clamp_chunk_params,
)


# ── the chunker's contract, which CHANGED on purpose ───────────────────────
#
# What used to be here: a byte-for-byte parity test against the original flat
# paragraph algorithm, guarding it against accidental change. That guard did its
# job and is now retired, because the change was not accidental — chunking became
# structure-aware (govern_doc_blocks), and the two cannot both be right:
#
#   * a chunk never straddles a heading, so section boundaries split
#   * heading text is NOT copied into the body; it lives in `heading_path`
#   * tables are whole blocks, never merged into prose
#   * `strategy` and `overlap` are accepted and ignored — structure is derived
#     from the document rather than guessed at by a parameter
#
# Keeping the old assertions would have meant keeping the old chunker beside the
# new one, which is the thing this refactor set out to remove.
def test_chunk_doc_returns_the_indexable_passages():
    chunks = chunk_doc("# A\n\nMot doan van.\n\n## B\n\nDoan van khac.")
    assert all(isinstance(c, str) and c.strip() for c in chunks)
    # Section boundaries split, so no chunk holds both sections' prose.
    assert not any("Mot doan van" in c and "Doan van khac" in c for c in chunks)


def test_heading_text_is_not_duplicated_into_the_body():
    """It lives in `heading_path`. A reader that shows a passage shows its heading
    path with it, and duplicating it would embed every heading twice."""
    chunks = chunk_doc("## Cam ket giao dung hen\n\nMuc tieu 92%.")
    assert any("Muc tieu 92%" in c for c in chunks)
    assert not any(c.startswith("## Cam ket") for c in chunks)


def test_strategy_and_overlap_are_accepted_and_ignored():
    """Every saved document still carries a chunk profile, so the kwargs cannot
    just disappear — but structure now comes from the document."""
    body = "# T\n\nDoan mot.\n\n## S\n\nDoan hai."
    assert chunk_doc(body) == chunk_doc(body, strategy="fixed", overlap=40)


def test_chunk_doc_heading_strategy_splits_on_headings():
    body = "# A\n\nintro to A\n\n## A.1\n\n" + ("detail " * 300) + "\n\n# B\n\nintro to B"
    chunks = chunk_doc(body, strategy="heading", size=200)
    assert len(chunks) >= 2
    # section A's content never bleeds into section B's chunk
    assert not any("intro to B" in c and "intro to A" in c for c in chunks)


def test_chunk_doc_fixed_strategy_ignores_paragraph_boundaries():
    body = "word " * 500
    chunks = chunk_doc(body, strategy="fixed", size=200)
    assert all(len(c) <= 200 + 20 for c in chunks)  # +overlap slack (0 here) is fine at 0


def test_preview_is_the_same_code_path_as_indexing():
    """`overlap` used to prepend the previous chunk's tail, and that assertion
    lived here. The block chunker derives boundaries from structure instead, so
    there is no tail to prepend — what still matters is that a PREVIEW cannot
    disagree with what will actually be embedded."""
    body = "# A\n\nMot doan.\n\n## B\n\n| c1 | c2 |\n|---|---|\n| 1 | 2 |\n\nDoan cuoi."
    assert preview_chunks(body) == chunk_doc(body)


def test_preview_chunks_matches_chunk_doc_exactly():
    body = "# Heading\n\nSome text.\n\n## Sub\n\nMore text " * 5
    for kwargs in ({}, {"strategy": "heading", "size": 300, "overlap": 15}, {"strategy": "fixed", "size": 100}):
        assert preview_chunks(body, **kwargs) == chunk_doc(body, **kwargs)


def test_clamp_chunk_params_bounds():
    assert _clamp_chunk_params(50, 0) == (100, 0)      # floor
    assert _clamp_chunk_params(5000, 0) == (1400, 0)   # ceiling (_HARD)
    size, overlap = _clamp_chunk_params(400, 900)
    assert overlap <= size // 2                         # overlap can't exceed half the chunk


# ── file_text_extractor round-trip on tiny real fixtures ────────────────────
def test_extract_text_docx_roundtrip():
    import docx
    from app.services.govern_doc_sources.file_text_extractor import extract_text

    document = docx.Document()
    document.add_paragraph("Hello from a test DOCX.")
    document.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    document.save(buf)

    result = extract_text(buf.getvalue(), "sample.docx")
    assert result["ok"] is True
    assert "Hello from a test DOCX." in result["text"]
    assert "Second paragraph." in result["text"]


def test_extract_text_xlsx_roundtrip():
    from openpyxl import Workbook
    from app.services.govern_doc_sources.file_text_extractor import extract_text

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "value"])
    ws.append(["gmv", 123])
    buf = io.BytesIO()
    wb.save(buf)

    result = extract_text(buf.getvalue(), "sample.xlsx")
    assert result["ok"] is True
    assert "gmv" in result["text"]
    assert "123" in result["text"]


def test_extract_text_pdf_roundtrip():
    from reportlab.pdfgen import canvas
    from app.services.govern_doc_sources.file_text_extractor import extract_text

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(100, 750, "Hello from a test PDF.")
    c.save()

    result = extract_text(buf.getvalue(), "sample.pdf")
    assert result["ok"] is True
    assert "Hello from a test PDF." in result["text"]


def test_extract_text_unsupported_extension():
    from app.services.govern_doc_sources.file_text_extractor import extract_text
    result = extract_text(b"whatever", "sample.txt")
    assert result["ok"] is False
    assert "Unsupported" in result["error"]


# ── google_doc_fetcher: pure text-flattening logic ──────────────────────────
def test_google_doc_flatten_headings_and_paragraphs():
    from app.services.govern_doc_sources.google_doc_fetcher import _flatten, extract_google_doc_id

    doc = {
        "body": {
            "content": [
                {"paragraph": {"paragraphStyle": {"namedStyleType": "HEADING_1"}, "elements": [{"textRun": {"content": "Title\n"}}]}},
                {"paragraph": {"paragraphStyle": {"namedStyleType": "NORMAL_TEXT"}, "elements": [{"textRun": {"content": "Body text.\n"}}]}},
                {"paragraph": {"paragraphStyle": {"namedStyleType": "HEADING_2"}, "elements": [{"textRun": {"content": "Sub\n"}}]}},
            ]
        }
    }
    text = _flatten(doc)
    assert text == "# Title\n\nBody text.\n\n## Sub"


def test_extract_google_doc_id_from_url_and_bare_id():
    from app.services.govern_doc_sources.google_doc_fetcher import extract_google_doc_id
    assert extract_google_doc_id("https://docs.google.com/document/d/1AbCdEfGhIjKlMnOpQr/edit") == "1AbCdEfGhIjKlMnOpQr"
    assert extract_google_doc_id("1AbCdEfGhIjKlMnOpQr") == "1AbCdEfGhIjKlMnOpQr"


# ── A · the chunk cap is a runaway guard, not a silent editorial cut ────────
def test_long_document_is_indexed_whole_and_reports_no_truncation():
    """A 38k-char document used to lose a third of itself to a 40-chunk cap
    with no signal anywhere. Coverage must now be effectively complete."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import chunk_doc_detailed

    body = "\n\n".join(f"## Muc {i}\n\n" + ("Noi dung doan van so %d. " % i) * 18 for i in range(1, 121))
    assert len(body) > 30_000
    chunks, stats = chunk_doc_detailed(body)

    assert stats["truncated"] is False
    assert stats["dropped_chunks"] == 0 and stats["dropped_chars"] == 0
    covered = sum(len(c) for c in chunks)
    # Not 100%, and not a defect: heading text is deliberately absent from the
    # bodies (it is carried in `heading_path`), so a document of 120 headings
    # loses those characters from this count while losing nothing from the index.
    assert covered >= len(body) * 0.95, f"only {covered}/{len(body)} chars indexed"


def test_truncation_is_reported_not_silent():
    """When the cap IS hit, the caller must be able to tell the user."""
    from app.services.dashboard_ai_bot import govern_doc_embeddings as gde

    # Deliberately far past the cap: 500 chunks x 100 chars only covers 50k, so
    # a ~92k body guarantees the guard actually fires.
    body = "\n\n".join(f"Doan van so {i} voi noi dung du dai de vuot tran." for i in range(1800))
    assert len(body) > 90_000
    chunks, stats = gde.chunk_doc_detailed(body, strategy="fixed", size=100)

    assert stats["truncated"] is True
    assert len(chunks) == stats["kept"] == gde._MAX_CHUNKS
    # `produced` must mean "what this document WOULD have produced", or
    # "truncated: true, produced: 500, kept: 500" is a report that contradicts
    # itself — which is what it said before this was fixed.
    assert stats["produced"] > gde._MAX_CHUNKS
    assert stats["produced"] == stats["kept"] + stats["dropped_chunks"]
    assert stats["dropped_chunks"] > 0 and stats["dropped_chars"] > 0


def test_chunk_doc_still_returns_a_plain_list():
    """chunk_doc() is the public API used across the codebase; adding stats must
    not change its return shape."""
    out = chunk_doc("Mot doan van ngan.")
    assert isinstance(out, list) and all(isinstance(c, str) for c in out)


# ── C · hybrid retrieval: fusion + accent-folded keyword SQL ────────────────
def test_rrf_fusion_lifts_the_item_both_rankers_agree_on():
    """RRF must reward agreement: an item ranked mid-table by both beats an
    item that only one ranker put first. That is the whole point of hybrid."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _fuse_rrf

    vector = ["a", "b", "c"]
    keyword = ["d", "b", "e"]
    scores = _fuse_rrf(vector, keyword)
    ranked = sorted(scores, key=lambda i: -scores[i])
    assert ranked[0] == "b", ranked


def test_rrf_uses_order_only_never_raw_scores():
    """Cosine and ts_rank are different scales; fusion must depend on position
    alone, so two identically-ordered lists produce identical scores."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _fuse_rrf

    assert _fuse_rrf(["x", "y"]) == _fuse_rrf(["x", "y"])
    assert _fuse_rrf(["x"])["x"] == pytest.approx(1.0 / 61)


def test_keyword_sql_folds_document_and_query_identically():
    """Folding only one side matches nothing at all — a silent, total failure."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _keyword_sql

    sql = _keyword_sql(" WHERE c.doc_id = :d ", "appbi_unaccent")
    assert "to_tsvector('simple', appbi_unaccent(c.content))" in sql
    assert "plainto_tsquery('simple', appbi_unaccent(:q))" in sql
    assert sql.count("appbi_unaccent") == 4  # both sides, in match AND in ts_rank

    plain = _keyword_sql(" WHERE c.doc_id = :d ", "")
    assert "appbi_unaccent" not in plain and "to_tsvector('simple', (c.content))" in plain


def test_keyword_ranking_never_breaks_vector_search():
    """Keyword ranking is an enhancement. If the DB cannot serve it, retrieval
    must degrade to vector-only rather than fail the user's question."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _keyword_ranked_ids

    class _DeadDb:
        def execute(self, *a, **k):
            raise RuntimeError("no such function: appbi_unaccent")

        def rollback(self):
            pass

    assert _keyword_ranked_ids(_DeadDb(), " WHERE 1=1 ", {}, "bat ky", 10) == []


# ── D · PDF tables must survive as tables ──────────────────────────────────
def _table_pdf() -> bytes:
    """A PDF whose only content is a ruled financial table."""
    reportlab = pytest.importorskip("reportlab")  # noqa: F841
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    from reportlab.lib import colors

    buf = io.BytesIO()
    rows = [
        ["Hang muc", "Muc tieu", "Thuc te", "Chenh lech"],
        ["Doanh thu", "100 ty", "92 ty", "-8 ty"],
        ["Chi phi", "60 ty", "58 ty", "-2 ty"],
    ]
    table = Table(rows)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]))
    SimpleDocTemplate(buf, pagesize=A4).build([table])
    return buf.getvalue()


def test_extract_pdf_keeps_table_rows_and_columns():
    """Flat text extraction turns this table into an unreadable number soup —
    'Doanh thu / 100 ty / 92 ty' with nothing saying which is target vs actual.
    A BI product cannot afford to embed that."""
    pytest.importorskip("pdfplumber")
    from app.services.govern_doc_sources.file_text_extractor import extract_text

    res = extract_text(_table_pdf(), "bao-cao.pdf")
    assert res["ok"], res.get("error")
    text = res["text"]

    assert "|" in text, f"table was flattened, not preserved:\n{text}"
    row = next((ln for ln in text.splitlines() if "Doanh thu" in ln), "")
    for cell in ("100 ty", "92 ty", "-8 ty"):
        assert cell in row, f"'{cell}' left its row — row was:\n{row!r}"


def test_extract_pdf_does_not_duplicate_table_cells_as_prose():
    """Rendering the table AND re-extracting its cells as loose text would embed
    every number twice and skew retrieval."""
    pytest.importorskip("pdfplumber")
    from app.services.govern_doc_sources.file_text_extractor import extract_text

    text = extract_text(_table_pdf(), "bao-cao.pdf")["text"]
    assert text.count("Doanh thu") == 1, text


# ── K1 · the approximate index must not be trusted for ORDER ────────────────
class _FakeDb:
    """Minimal Session stand-in: records SET statements, replays canned rows."""

    def __init__(self, rows=(), fail_on=None):
        self.rows = list(rows)
        self.fail_on = fail_on or ()
        self.statements = []
        self.rolled_back = 0

    def execute(self, stmt, params=None):
        sql = str(stmt)
        self.statements.append(sql)
        if any(token in sql for token in self.fail_on):
            raise RuntimeError("unsupported")
        return self

    def fetchall(self):
        return self.rows

    def rollback(self):
        self.rolled_back += 1


def test_vector_ranked_ids_sorts_by_distance_not_index_order():
    """relaxed_order buys recall by letting the index emit rows slightly out of
    order. That is only safe because we re-sort on the distance we selected —
    RRF fuses by RANK, so a wrong order corrupts fusion invisibly."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _vector_ranked_ids

    # Index hands them back scrambled; nearest is id 7 (0.05).
    db = _FakeDb(rows=[(3, 0.42), (7, 0.05), (9, 0.31)])
    assert _vector_ranked_ids(db, " WHERE 1=1 ", {}, "[0,0]", 10) == [7, 9, 3]


def test_vector_scan_is_deterministic():
    """`relaxed_order` returns rows as it finds them, so the candidate SET could
    differ between two runs of byte-identical SQL — the eval harness caught the
    dashboard and agent suites disagreeing intermittently, about one run in three.
    A knowledge base that answers the same question differently on a re-ask cannot
    be audited, and at ef_search 400 strict ordering costs no measurable recall."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import (
        _HNSW_EF_SEARCH, _HNSW_ITERATIVE_SCAN,
    )

    assert _HNSW_ITERATIVE_SCAN == "strict_order"
    assert _HNSW_EF_SEARCH >= 400, "strict ordering needs the wider search to keep recall"


def test_vector_scan_tuning_is_applied_before_the_query():
    from app.services.dashboard_ai_bot.govern_doc_embeddings import (
        _HNSW_EF_SEARCH, _HNSW_ITERATIVE_SCAN, _vector_ranked_ids,
    )

    db = _FakeDb(rows=[(1, 0.1)])
    _vector_ranked_ids(db, " WHERE 1=1 ", {}, "[0,0]", 5)
    joined = "\n".join(db.statements)
    assert f"hnsw.ef_search = {_HNSW_EF_SEARCH}" in joined
    assert f"hnsw.iterative_scan = {_HNSW_ITERATIVE_SCAN}" in joined
    # Tuning must come FIRST — a SET after the scan tunes nothing.
    assert db.statements.index([s for s in db.statements if "ef_search" in s][0]) < len(db.statements) - 1


def test_vector_scan_tuning_failure_does_not_break_retrieval():
    """A server without HNSW must still answer questions."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _tune_vector_scan

    db = _FakeDb(fail_on=("hnsw.",))
    _tune_vector_scan(db)              # must not raise
    assert db.rolled_back == 1         # and must clear the poisoned transaction


# ── K2 · one index, one model ──────────────────────────────────────────────
def test_configured_embedding_profiles_share_the_pgvector_dimensions():
    from app.core.config import settings
    from app.services.embedding_service import EmbeddingService

    profiles = EmbeddingService.embedding_profiles()
    assert {p["model"] for p in profiles} >= {
        "text-embedding-3-small", "text-embedding-3-large",
    }
    assert {p["dimensions"] for p in profiles} == {
        settings.openai_embedding_dimensions
    }


def test_unknown_embedding_model_is_rejected_before_indexing():
    from app.services.embedding_service import EmbeddingService

    with pytest.raises(ValueError, match="Unsupported embedding model"):
        EmbeddingService.resolve_model("text-embedding-ada-002")


def test_model_mismatch_is_relative_to_the_document_not_a_global_default():
    from app.services.governance_service import GovernanceService

    class _Doc:
        id = 1
        embedding_model = "text-embedding-3-large"

    def check(models):
        return GovernanceService._model_mismatched(_FakeDb(rows=[(m,) for m in models]), _Doc())

    assert check(["text-embedding-3-large"]) is False
    assert check(["text-embedding-3-small"]) is True
    assert check(["text-embedding-3-large", "text-embedding-3-small"]) is True
    assert check([]) is False


def test_reset_switches_model_only_after_deleting_old_vectors(monkeypatch):
    from app.services.dashboard_ai_bot import govern_doc_embeddings as gde

    class _Result:
        def fetchall(self):
            return []

    class _Db:
        def __init__(self):
            self.statements = []
            self.commits = 0

        def execute(self, stmt, params=None):
            self.statements.append((str(stmt), params or {}))
            return _Result()

        def commit(self):
            self.commits += 1

        def rollback(self):
            pass

        def refresh(self, _doc):
            pass

        def flush(self):
            pass

    class _Doc:
        id = 9
        embedded_hash = "old"
        embedding_model = "text-embedding-3-small"
        chunk_strategy = "paragraph"
        chunk_size = 850
        chunk_overlap = 0

    # Reset persists the new profile and QUEUES the rebuild. It used to embed
    # inline, which made changing a chunking profile — the single largest bulk
    # operation in the module — hold one HTTP request open for the whole document.
    queued = []
    import app.services.govern_doc_index_queue as queue_mod

    monkeypatch.setattr(
        queue_mod,
        "enqueue",
        lambda db, doc_id, **kwargs: queued.append((doc_id, kwargs.get("reason")))
        or {"id": 1, "state": "queued", "reason": kwargs.get("reason")},
    )
    db, doc = _Db(), _Doc()
    result = gde.reset_doc_embedding(
        db,
        doc,
        model="text-embedding-3-large",
        chunk_strategy="paragraph",
        chunk_size=850,
        chunk_overlap=0,
        changed_by="tester",
    )

    statements = " ".join(text for text, _params in db.statements)
    # Old vectors are deleted and the row is locked BEFORE the profile changes, so
    # a provider failure can never leave chunks labelled with a model that did not
    # produce them.
    assert "FOR UPDATE" in statements
    assert "DELETE FROM govern_doc_chunk" in statements
    assert doc.embedding_model == "text-embedding-3-large"
    assert doc.embedded_hash is None

    assert queued == [(9, "config")]
    assert result["status"] == "queued"
    assert result["embedding_model"] == "text-embedding-3-large"


def test_multi_model_search_embeds_the_query_once_per_model(monkeypatch):
    from app.services.dashboard_ai_bot import govern_doc_embeddings as gde
    from app.services.embedding_service import EmbeddingService

    class _Rows:
        # 12 columns now: the SELECT carries the citation (heading path, page,
        # block kind) and the section key, because a passage that cannot say where
        # it came from cannot be cited.
        def fetchall(self):
            return [
                chunk_row(id=10, doc_id=1, title="Doc A", content="alpha",
                          model_version="model-a", heading_path="Doc A"),
                chunk_row(id=20, doc_id=2, title="Doc B", content="beta",
                          model_version="model-b", heading_path="Doc B"),
            ]

        def first(self):
            return (2, 100.0)

    class _Db:
        # The retriever now also asks for corpus statistics (for BM25) and for the
        # surrounding section. Asserting on EVERY statement would make this a test
        # of the SQL inventory rather than of the per-model query behaviour.
        def execute(self, stmt, params=None):
            return _Rows()

        def rollback(self):
            pass

    generated = []
    monkeypatch.setattr(
        gde, "_model_doc_groups", lambda *_args: {"model-a": [1], "model-b": [2]}
    )
    monkeypatch.setattr(gde, "_keyword_ranked_ids", lambda *_args: [])
    monkeypatch.setattr(
        EmbeddingService,
        "generate_query_embedding",
        lambda query, model=None: generated.append((query, model)) or [0.1, 0.2],
    )
    monkeypatch.setattr(
        gde,
        "_vector_ranked_hits",
        lambda _db, _scope, params, _vector, _limit: [
            (10, 0.91) if params["embedding_model"] == "model-a" else (20, 0.82)
        ],
    )

    rows = gde._search_scoped_doc_chunks(
        _Db(), "revenue policy", k=5, dashboard_id=None,
        doc_ids={1, 2}, published_only=True,
    )
    assert generated == [
        ("revenue policy", "model-a"),
        ("revenue policy", "model-b"),
    ]
    assert {row["chunk_id"] for row in rows} == {10, 20}
    assert {row["embedding_model"] for row in rows} == {"model-a", "model-b"}


def test_failed_model_keeps_its_keyword_results(monkeypatch):
    from app.services.dashboard_ai_bot import govern_doc_embeddings as gde
    from app.services.embedding_service import EmbeddingService

    class _Rows:
        def fetchall(self):
            return [chunk_row(id=20, doc_id=2, title="Doc B",
                              content="exact code Q2", model_version="model-b",
                              heading_path="Doc B > Muc", token_count=12)]

    class _Db:
        def execute(self, stmt, params=None):
            return _Rows()

    monkeypatch.setattr(gde, "_model_doc_groups", lambda *_args: {"model-b": [2]})
    monkeypatch.setattr(gde, "_keyword_ranked_ids", lambda *_args: [20])
    monkeypatch.setattr(EmbeddingService, "generate_query_embedding", lambda *_args, **_kwargs: None)

    rows = gde._search_scoped_doc_chunks(
        _Db(), "Q2", k=5, dashboard_id=None,
        doc_ids={2}, published_only=True,
    )
    assert len(rows) == 1
    assert rows[0]["chunk_id"] == 20
    assert rows[0]["matched_by"] == "keyword"
    assert rows[0]["similarity"] is None


# ── S1 · the authoring window must be opened AND closed explicitly ──────────
def test_authoring_scope_opens_and_restricted_scope_closes():
    """SET LOCAL survives to the end of the transaction, so an authoring read
    earlier in a request would leave drafts visible to everything after it.
    Retrieval must therefore close the window rather than assume it is shut."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import authoring_scope, restricted_scope

    db = _FakeDb()
    authoring_scope(db)
    restricted_scope(db)
    assert all("SET LOCAL appbi.chunk_scope" in s for s in db.statements)
    assert len(db.statements) == 2


def test_chunk_scope_is_bound_not_interpolated():
    """The scope value reaches Postgres as a parameter, so it can never carry SQL."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import authoring_scope

    db = _FakeDb()
    authoring_scope(db)
    assert ":v" in db.statements[0] and "authoring" not in db.statements[0]


def test_chunk_scope_failure_clears_the_transaction():
    """A server without the GUC must not leave a poisoned transaction behind."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import restricted_scope

    db = _FakeDb(fail_on=("appbi.chunk_scope",))
    restricted_scope(db)
    assert db.rolled_back == 1


def test_retrieval_closes_the_authoring_window_before_reading():
    """Guards the ORDER, which is the whole point: closing the scope after the
    query would protect nothing."""
    import inspect

    from app.services.dashboard_ai_bot import govern_doc_embeddings as gde

    src = inspect.getsource(gde.search_doc_chunks)
    assert "restricted_scope(db)" in src
    assert src.index("restricted_scope(db)") < src.index("_search_scoped_doc_chunks(")


# ── S4 · nothing leaves without a decision, and every transfer is recorded ──
def test_an_unset_policy_defaults_to_the_column_default():
    """A fresh ORM object has the field as None until flush. Treating that as the
    most restrictive value once left every brand-new document unindexed."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import (
        egress_allowed, processing_policy,
    )

    class D:
        pass

    d = D()
    assert processing_policy(d) == "embedding"       # attribute absent entirely
    d.external_processing = None                     # constructed, not flushed
    assert processing_policy(d) == "embedding"
    d.external_processing = "nonsense"               # unrecognised, not trusted
    assert processing_policy(d) == "embedding"
    assert egress_allowed(d, "embedding") is True


def test_each_purpose_requires_its_own_level():
    """The boolean was named after ONE outbound call. OCR and figure description
    send page IMAGES, so a document marked "text only" must refuse them — that is
    the whole reason the flag became a policy."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import egress_allowed

    class D:
        def __init__(self, level):
            self.external_processing = level

    assert egress_allowed(D("embedding"), "embedding") is True
    assert egress_allowed(D("embedding"), "ocr") is False
    assert egress_allowed(D("embedding"), "vision") is False

    assert egress_allowed(D("full"), "embedding") is True
    assert egress_allowed(D("full"), "ocr") is True

    for purpose in ("embedding", "ocr", "vision", "rerank"):
        assert egress_allowed(D("none"), purpose) is False

    # An unknown purpose is refused unless the policy is fully open, so adding a
    # new external call cannot accidentally inherit permission.
    assert egress_allowed(D("embedding"), "something_new") is False


def test_absent_policy_key_never_flips_the_policy():
    """The API builds the payload with model_dump(), so every optional key is
    PRESENT with value None. Keying on presence turned an ordinary metadata save
    into 'block this document'."""
    import inspect

    from app.services.governance_service import GovernanceService

    src = inspect.getsource(GovernanceService.upsert_knowledge_doc)
    assert 'payload.get("external_processing") is not None' in src
    assert '"external_processing" in payload' not in src


def test_egress_log_write_failure_never_breaks_indexing():
    from app.services.dashboard_ai_bot.govern_doc_embeddings import log_egress

    class D:
        id = 1
        title = "x"
        sensitivity = "internal"

    db = _FakeDb(fail_on=("govern_doc_egress_log",))
    log_egress(db, D(), outcome="sent", model="m", chunks=1, chars=10)   # must not raise
    assert db.rolled_back == 1


# ── Retrieval engine · keyword branch and reproducibility ───────────────────
def test_keyword_branch_matches_on_or_not_and():
    """`plainto_tsquery` joins tokens with AND, so a real question became a
    twelve-term AND chain that almost no passage satisfied — measured: the keyword
    branch reached 3 of 23 questions. Relaxing to OR took recall@1 from 0.60 to
    0.85. This guards the operator, which is easy to lose in a refactor because
    the AND form reads like the more obvious code."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _keyword_sql

    sql = _keyword_sql(" WHERE c.doc_id = :d ", "appbi_unaccent")
    assert "' & ', ' | '" in sql, "the AND -> OR rewrite is gone"
    assert "::tsquery" in sql


def test_ranking_sql_is_totally_ordered():
    """Ties became common once the keyword query was relaxed to OR, and RRF fuses
    by RANK — so without a total order the fused ranking is not reproducible. The
    eval harness caught the dashboard and agent paths disagreeing on a case with
    identical inputs, from tie order alone."""
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _keyword_sql

    assert "DESC, c.id" in _keyword_sql(" WHERE 1=1 ", "appbi_unaccent")


def test_vector_hits_break_ties_deterministically():
    from app.services.dashboard_ai_bot.govern_doc_embeddings import _vector_ranked_hits

    # Two chunks at the SAME distance, handed back in an unhelpful order.
    db = _FakeDb(rows=[(9, 0.25), (4, 0.25), (7, 0.10)])
    assert [chunk_id for chunk_id, _ in _vector_ranked_hits(db, " WHERE 1=1 ", {}, "[0,0]", 10)] \
        == [7, 4, 9]


def test_stale_index_scan_can_include_drafts():
    """The authoring console searches drafts, so a scan that only looks at
    Published documents has no detector for the case it actually broke on: doc 43
    had a NULL hash, the retriever refused every chunk, and the Vectors tab showed
    an empty result with nothing saying why."""
    import inspect

    from app.services.dashboard_ai_bot import govern_doc_embeddings as gde

    signature = inspect.signature(gde.stale_index_docs)
    assert "published_only" in signature.parameters
    assert signature.parameters["published_only"].default is True
